"""
Utility helpers for exporting generated documents to DOCX and PDF.

This module is imported lazily by the FastAPI export endpoint.  We keep the
dependencies lightweight and rely only on packages that already ship with the
backend stack (python-docx, Pillow, pypdf).
"""

from __future__ import annotations

import io
import re
import textwrap
from dataclasses import dataclass
from html import unescape
from typing import Iterable, List, Tuple

from docx import Document  # type: ignore
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore
from docx.shared import Pt  # type: ignore
from PIL import Image, ImageDraw, ImageFont  # type: ignore


__all__ = [
    "export_document",
]


DEFAULT_DPI = 300
A4_WIDTH_PX = int(round(8.27 * DEFAULT_DPI))   # ≈ 210mm
A4_HEIGHT_PX = int(round(11.69 * DEFAULT_DPI))  # ≈ 297mm
PAGE_MARGIN_PX = int(round(0.7 * DEFAULT_DPI))  # ≈ 18mm
DEFAULT_FONT = "DejaVuSans.ttf"
DEFAULT_FONT_POINT_SIZE = 11
DEFAULT_FONT_SIZE = max(24, int(round(DEFAULT_FONT_POINT_SIZE * DEFAULT_DPI / 72)))


@dataclass
class ExportResult:
    data: bytes
    filename: str
    media_type: str


def export_document(*, format: str, title: str, content: str) -> Tuple[bytes, str, str]:
    """
    Export a document in the requested format.

    :param format: Requested export format ("pdf" | "docx").
    :param title: Document title used for heading and filename.
    :param content: HTML content of the generated document.
    :return: (bytes, filename, media_type)
    """
    normalized_format = (format or "docx").lower()
    title = title.strip() or "Rechtsdokument"
    sanitized_content = sanitize_html(content)

    if normalized_format == "pdf":
        result = export_pdf(title=title, text=sanitized_content)
    elif normalized_format == "docx":
        result = export_docx(title=title, text=sanitized_content)
    else:
        raise ValueError(f"Unsupported export format: {format}")

    return result.data, result.filename, result.media_type


def sanitize_html(html: str | None) -> str:
    """
    Reduce HTML from the editor into plain text while keeping a readable layout.
    """
    if not html:
        return ""

    text = html
    # Normalize line breaks for block-level elements.
    replacements = {
        r"(?i)<\s*/\s*p\s*>": "\n\n",
        r"(?i)<\s*br\s*/?\s*>": "\n",
        r"(?i)<\s*/\s*li\s*>": "\n",
        r"(?i)<\s*h[1-6][^>]*>": "\n\n",
        r"(?i)<\s*/\s*h[1-6]\s*>": "\n\n",
    }
    for pattern, value in replacements.items():
        text = re.sub(pattern, value, text)

    # Replace list items with bullets.
    text = re.sub(r"(?is)<\s*li[^>]*>", "• ", text)

    # Strip all remaining tags.
    text = re.sub(r"<[^>]+>", "", text)

    # Decode HTML entities and collapse whitespace.
    text = unescape(text)
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"\u00a0", " ", text)  # non-breaking space
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()

    return text


def slugify(value: str, fallback: str = "dokument") -> str:
    """
    Create a filename-friendly slug from the given string.
    """
    value = value.lower()
    value = re.sub(r"[^a-z0-9\s-]", "", value)
    value = re.sub(r"[\s_-]+", "-", value).strip("-")
    return value or fallback


def export_docx(*, title: str, text: str) -> ExportResult:
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    # Title heading
    heading = document.add_heading(title, level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for paragraph in split_paragraphs(text):
        p = document.add_paragraph(paragraph)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)

    buffer = io.BytesIO()
    document.save(buffer)
    filename = f"{slugify(title)}.docx"
    return ExportResult(
        data=buffer.getvalue(),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def export_pdf(*, title: str, text: str) -> ExportResult:
    font = load_font(DEFAULT_FONT_SIZE)
    pages = render_text_to_pages(title=title, text=text, font=font)

    if not pages:
        pages = [Image.new("RGB", (A4_WIDTH_PX, A4_HEIGHT_PX), "white")]

    buffer = io.BytesIO()
    first, *rest = pages
    first.save(buffer, format="PDF", save_all=bool(rest), append_images=rest)
    filename = f"{slugify(title)}.pdf"
    return ExportResult(data=buffer.getvalue(), filename=filename, media_type="application/pdf")


def split_paragraphs(text: str) -> List[str]:
    if not text:
        return []
    return [segment.strip() for segment in text.split("\n\n") if segment.strip()]


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    """
    Load a truetype font if available, otherwise fall back to the default bitmap font.
    """
    try:
        return ImageFont.truetype(DEFAULT_FONT, size=size)
    except Exception:
        try:
            return ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", size=size)
        except Exception:
            return ImageFont.load_default()


def render_text_to_pages(*, title: str, text: str, font: ImageFont.ImageFont) -> List[Image.Image]:
    wrapped_lines = wrap_text_for_pdf(title=title, text=text, font=font)
    line_height = max(int(_line_height(font)), int(DEFAULT_FONT_SIZE * 1.2))

    max_lines_per_page = (A4_HEIGHT_PX - 2 * PAGE_MARGIN_PX) // line_height
    pages: List[Image.Image] = []

    for page_index in range(0, len(wrapped_lines), max_lines_per_page):
        chunk = wrapped_lines[page_index : page_index + max_lines_per_page]
        img = Image.new("RGB", (A4_WIDTH_PX, A4_HEIGHT_PX), "white")
        draw = ImageDraw.Draw(img)
        y = PAGE_MARGIN_PX
        for line in chunk:
            draw.text((PAGE_MARGIN_PX, y), line, fill="black", font=font)
            y += line_height
        pages.append(img)

    return pages


def wrap_text_for_pdf(*, title: str, text: str, font: ImageFont.ImageFont) -> List[str]:
    content_lines: List[str] = []
    available_width = A4_WIDTH_PX - 2 * PAGE_MARGIN_PX

    def _wrap_line(line: str) -> Iterable[str]:
        if not line:
            return [""]
        # Use binary search with textwrap for predictable wrapping.
        wrapped: List[str] = []
        for raw in line.split("\n"):
            if not raw.strip():
                wrapped.append("")
                continue
            words = raw.split()
            current = []
            for word in words:
                test_line = " ".join(current + [word])
                if _text_length(font, test_line) <= available_width:
                    current.append(word)
                else:
                    if not current:
                        # In case of a single very long word, fall back to textwrap.
                        for segment in textwrap.wrap(word, width=40):
                            wrapped.append(segment)
                        current = []
                    else:
                        wrapped.append(" ".join(current))
                        current = [word]
            if current:
                wrapped.append(" ".join(current))
        return wrapped or [""]

    # Title as first line with extra spacing.
    if title:
        content_lines.extend(_wrap_line(title))
        content_lines.append("")

    for paragraph in split_paragraphs(text):
        content_lines.extend(_wrap_line(paragraph))
        content_lines.append("")

    # Remove trailing blank lines.
    while content_lines and not content_lines[-1].strip():
        content_lines.pop()

    return content_lines or [title]


def _text_length(font: ImageFont.ImageFont, text: str) -> float:
    if hasattr(font, "getlength"):
        try:
            return font.getlength(text)
        except Exception:
            pass
    if hasattr(font, "getbbox"):
        try:
            bbox = font.getbbox(text)
            return max(bbox[2] - bbox[0], 0)
        except Exception:
            pass
    try:
        size = font.getsize(text)
        return max(size[0], 0)
    except Exception:
        return float(len(text) * DEFAULT_FONT_SIZE)


def _line_height(font: ImageFont.ImageFont) -> float:
    if hasattr(font, "getmetrics"):
        try:
            ascent, descent = font.getmetrics()
            return ascent + descent + 6
        except Exception:
            pass
    if hasattr(font, "getbbox"):
        try:
            bbox = font.getbbox("Ag")
            return bbox[3] - bbox[1]
        except Exception:
            pass
    try:
        size = font.getsize("Ag")
        return size[1]
    except Exception:
        return DEFAULT_FONT_SIZE * 1.2
